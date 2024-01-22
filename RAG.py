from pypdf import PdfReader
import numpy as np
import pandas as pd
from tqdm import tqdm
import docx
from docx import Document
from dataclasses import dataclass
from enum import Enum, auto
from typing import List, Optional, Union, Tuple
from pathlib import Path
from glob import glob
import os
import re
from transformers import AutoTokenizer, AutoModel
import torch
# import pdfplumber

# 对文本进行拆分
CHUNK_SIZE = 256
# CHUNK_SIZE = 512
# global_dir = "政策归档文件"


class FileType(Enum):
    PDF = auto()
    doc = auto()
    docx = auto()


@dataclass
class TransOutput:
    file_name: str
    file_type: FileType
    text_data: Union[pd.DataFrame, None]


def transpdf(pdf_path: str, show_progress_bar: bool = False):
    print("[transpdf1]===============================")
    reader = PdfReader(pdf_path)
    number_of_pages = len(reader.pages)

    def page2text(pageid: int):
        page = reader.pages[pageid]
        text = page.extract_text()
        return text

    data = pd.DataFrame({
        # 'file_Path':pdf_path,
        'text': [page2text(i) for i in tqdm(range(number_of_pages), disable=not show_progress_bar)],
        'pageid': range(number_of_pages),
    })
    if type(pdf_path) != str:
        pdf_path = pdf_path.name
        
    res = TransOutput(
        file_name=pdf_path,
        file_type=FileType.PDF,
        text_data=data
    )
    print("---transpdf1:", res)
    return res

# def transpdf2(pdf_path: str, show_progress_bar: bool = False):
#     print("[transpdf2]===============================")
#     reader = pdfplumber.open(pdf_path)
#     # reader = PdfReader(pdf_path)
#     number_of_pages = len(reader.pages)

#     def page2text(pageid: int):
#         page = reader.pages[pageid]
#         text = page.extract_text()
#         return text

#     data = pd.DataFrame({
#         # 'file_Path':pdf_path,
#         'text': [page2text(i) for i in tqdm(range(number_of_pages), disable=not show_progress_bar)],
#         'pageid': range(number_of_pages),
#     })
#     res = TransOutput(
#         file_name=pdf_path.name,
#         file_type=FileType.PDF,
#         text_data=data
#     )
#     print("---transpdf2:", res)
#     return res

def transdocx(doc_path: str, show_progress_bar: bool = False):
    doc = docx.Document(doc_path)
    all_paras = doc.paragraphs
    number_of_pages = len(all_paras)
    
    # print("all_paras:", all_paras)
    # print("number_of_pages:", number_of_pages)
    
    raw_text = [i.text for i in tqdm(all_paras, disable=not show_progress_bar)]
    tmp_text = ""
    mid_text = []
    for idx, text in enumerate(raw_text):
        if len(tmp_text) + len(text) < CHUNK_SIZE:
            tmp_text += text
        else:
            tmp_text += " " * (CHUNK_SIZE - len(tmp_text))
            mid_text.append(tmp_text)
            tmp_text = ""
    
    if tmp_text != "":
        tmp_text += " " * (CHUNK_SIZE - len(tmp_text))
        mid_text.append(tmp_text)
        
    for tt in mid_text:
        print("debug_83:", len(tt), tt)

    data = pd.DataFrame({
        "text": mid_text,
        "paraid": range(len(mid_text))
    })
    # data = pd.DataFrame({
    #     'text': [i.text for i in tqdm(all_paras, disable=not show_progress_bar)],
    #     'paraid': range(number_of_pages)
    # })
    tp = ""
    if type(doc_path) == str:
        tp = doc_path
    else:
        tp = doc_path.name
        
    res = TransOutput(
            file_name=tp,
            file_type=FileType.docx,
            text_data=data
    )
    return res

def cal_detail_in_dir(dir_name):
    all_file_list = []
    # all_file_size = []

    for (root, dir, file_name) in os.walk(dir_name):
        for temp_file in file_name:
            standard_path = f"{root}/{temp_file}"

            all_file_list.append(standard_path)

    return all_file_list


def transfile(x: Path) -> TransOutput:
    if x.suffix == ".docx" :
        return transdocx(x.__str__())
    # elif x.suffix == ".doc":
        # return transdoc(x.__str__())
    elif x.suffix == ".pdf":
        return transpdf(x.__str__())
    else:
        print("unsupport file type")
        exit()


def cleanquestion(x: str) -> str:
    if isinstance(x, str):

        str_text = re.sub(
            u"([^\u4e00-\u9fa5\u0030-\u0039\u0041-\u005a\u0061-\u007a])", "", x)
        return str_text
    else:
        return None


def clean_text_data(transout: TransOutput) -> TransOutput:
    text_df = transout.text_data
    res = text_df.pipe(
        lambda x: x.assign(**{
            'new_text_': x['text'].apply(lambda j: cleanquestion(j))
        })
    ).pipe(
        lambda x: x.loc[x['new_text_'].apply(lambda j: len(j) > 0)]
    )

    transout.text_data = res
    return transout


def chunk_text(x: str) -> Union[None, List[str]]:
    if not isinstance(x, str):
        x = str(x)

    x_list = [x[startid:(startid + CHUNK_SIZE)] for startid in range(0, len(x), CHUNK_SIZE)]
    return x_list


def chunk_text4TransOutput(x: TransOutput) -> TransOutput:
    # try:
    text_df = x.text_data
    res = text_df.pipe(
        lambda x: x.assign(**{
            'chunk_text': x['new_text_'].apply(lambda j: chunk_text(j))
        })
    ).explode(['chunk_text']).drop(columns=['new_text_'])
    x.text_data = res
    return x
    # except Exception as e:
    #     return None


def numpy_cos_sim(a: np.ndarray, b: np.ndarray) -> np.ndarray:
    if len(a.shape) == 1:
        a = a.reshape(1, -1)
    if len(b.shape) == 1:
        b = b.reshape(1, -1)

    a_norm = a / np.linalg.norm(a, ord=2, axis=1).reshape(-1, 1)
    b_norm = b / np.linalg.norm(b, ord=2, axis=1).reshape(-1, 1)

    return np.matmul(a_norm, b_norm.T)


class SentenceVector:
    def __init__(self,
                 model_name_or_path: str = None,
                 device: str = "cuda:0") -> None:
        self.model_name_or_path = model_name_or_path
        self.device = device

        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name_or_path)

        self.model = AutoModel.from_pretrained(self.model_name_or_path)
        self.model.to(self.device)

    def encode_fun(self, texts: List[str]) -> np.ndarray:
        texts = [cleanquestion(i) for i in texts]

        inputs = self.tokenizer.batch_encode_plus(
            texts, padding=True, truncation=True, return_tensors="pt", max_length=64)
        inputs.to(device=self.device)
        with torch.no_grad():
            embeddings = self.model(**inputs)

        embeddings = embeddings.last_hidden_state[:, 0]
        embeddings = embeddings.to('cpu').numpy()
        return embeddings

    def encode_fun_plus(self, texts: List[str], batch_size: int = 100) -> np.ndarray:
        embeddings = np.concatenate([self.encode_fun(
            texts[i:(i + batch_size)]) for i in tqdm(range(0, len(texts), batch_size))])
        return embeddings


class KnowLedge:
    def __init__(self,
                 global_dir: str = None,
                 gen_model_name_or_path: str = "THUDM/chatglm-6b",
                 sen_embedding_model_name_or_path: str = "hfl/chinese-roberta-wwm-ext",
                 batch_top_k=5
                 ) -> None:

        self.batch_top_k = batch_top_k

        if global_dir is not None:

            all_file_list = cal_detail_in_dir(global_dir)
            all_file_list = [Path(i) for i in all_file_list]
            all_file_list = [i for i in all_file_list if i.suffix in ['.pdf', '.docx']]
            all_trans_data = [transfile(i) for i in tqdm(all_file_list)]
            all_trans_data = [clean_text_data(i) for i in all_trans_data]
            all_trans_data = [i for i in all_trans_data if i.text_data.shape[0] > 0]
            # print("all_trans_data:\n", all_trans_data)
            all_trans_data = [chunk_text4TransOutput(i) for i in all_trans_data]
        else:
            all_file_list = []
            all_trans_data = []

        self.sv = SentenceVector(model_name_or_path=sen_embedding_model_name_or_path)
        all_vector = [self.sv.encode_fun_plus(i.text_data['chunk_text'].tolist()) for i in all_trans_data]
        self.all_trans_data = all_trans_data
        self.all_vector = all_vector
        

        self.gen_tokenizer = AutoTokenizer.from_pretrained(gen_model_name_or_path, trust_remote_code=True)
        # self.gen_model = AutoModel.from_pretrained(gen_model_name_or_path, trust_remote_code=True, torch_dtype=torch.float16, device_map=auto)
        self.gen_model = AutoModel.from_pretrained(gen_model_name_or_path, trust_remote_code=True, load_in_8bit=True)

    def search_top_info(self, index: int, question_vector: np.ndarray) -> pd.DataFrame:
        # print("".format(index))
        similar_score = numpy_cos_sim(self.all_vector[index], question_vector).flatten()

        if similar_score.shape[0] < self.batch_top_k:
            res = self.all_trans_data[index].text_data.reset_index(drop=True).pipe(
                lambda x: x.assign(**{
                    'score': similar_score
                })
            ).pipe(
                lambda x: x.assign(**{
                    'file_name': self.all_trans_data[index].file_name,
                    'file_path': self.all_trans_data[index].file_type
                })
            )

        else:

            top_k_location = np.argpartition(similar_score, kth=-self.batch_top_k)[-self.batch_top_k:]

            res = self.all_trans_data[index].text_data.reset_index(drop=True).iloc[top_k_location].pipe(
                lambda x: x.assign(**{
                    'score': similar_score[top_k_location]
                })
            ).pipe(
                lambda x: x.assign(**{
                    'file_name': self.all_trans_data[index].file_name,
                    'file_path': self.all_trans_data[index].file_type
                })
            )

        return res

    # def search_result(self, question_str: str) -> Tuple[str, pd.DataFrame]:
    def search_result(self, question_str: str):

        # question_str = "做集成电路的企业,有什么补贴"#
        question_vector = self.sv.encode_fun([question_str])
        # question_vector.shape
        # index = 0
        print("debug_253: ", range(len(self.all_vector)))
        
        search_table_info = pd.concat(
            [self.search_top_info(index, question_vector) for index in range(len(self.all_vector))]).pipe(
            lambda x: x.sort_values(by=['score'], ascending=False)
        )
            
        search_table = search_table_info.drop_duplicates(['chunk_text']).head(30)
        print("debug_268:", search_table)

        search_text_list = search_table['chunk_text'].tolist()
        # len(search_text_list), search_text_list[:3]

        prompt_template = """基于以下已知信息，简洁和专业的来回答用户的问题。
        如果无法从中得到答案，请说 "根据已知信息无法回答该问题" 或 "没有提供足够的相关信息"，不允许在答案中添加编造成分，答案请使用中文。
        问题:
        {question}
        已知内容:
        {context}

        """

        text2chatglm = prompt_template.format_map({
            'question': question_str,
            'context': '\n'.join(search_text_list)
        })

        print("-------debug_319:")
        print(text2chatglm)
        self.text2chatglm = text2chatglm
        return search_table
        # # response = self.gen_model.generate(text2chatglm)
        # try:
        #     response, history = self.gen_model.chat(self.gen_tokenizer, text2chatglm, history=[])
        # except:
        #     response = "请重试"
        #     search_table = []
            
        # # self.gen_model = AutoModel.from_pretrained(gen_model_name_or_path, trust_remote_code=True, load_in_8bit=True)
        # torch.cuda.empty_cache()

        # return response, search_table
    
    def stream_search(self):
        return self.gen_model.stream_chat(self.gen_tokenizer, self.text2chatglm)
    
    def tel_QA(self, question_str: str, tel_str: str):
        # doc = docx.Document(file_str)
        # all_paras = doc.paragraphs
    
        # raw_text = [i.text for i in tqdm(all_paras)]
        prompt_template = """基于以下已知信息，简洁和专业的来回答用户的问题。
        如果无法从中得到答案，请说 "根据已知信息无法回答该问题" 或 "没有提供足够的相关信息"，不允许在答案中添加编造成分，答案请使用中文。
        问题:
        {question}
        已知内容:
        {context}

        """

        text2chatglm = prompt_template.format_map({
            'question': question_str,
            'context': tel_str
            # 'context': '\n'.join(raw_text)
        })

        response, history = self.gen_model.chat(self.gen_tokenizer, text2chatglm, history=[])
        torch.cuda.empty_cache()

        return response
 
        
    def reset_file(self, file_str: str, file_type: str):
        if file_type == "docx":
            all_trans_data = [transdocx(file_str)]
        else:
            all_trans_data = [transpdf(file_str)]
        all_trans_data = [clean_text_data(i) for i in all_trans_data]
        all_trans_data = [i for i in all_trans_data if i.text_data.shape[0] > 0]
        print("all_trans_data:\n", all_trans_data)

        all_trans_data = [chunk_text4TransOutput(i) for i in all_trans_data]

        all_vector = [self.sv.encode_fun_plus(i.text_data['chunk_text'].tolist()) for i in all_trans_data]

        self.all_trans_data = all_trans_data
        self.all_vector = all_vector
        
        print("------update file done")


    def reset_folder(self, global_dir: str):
        print("----in folder upload, ", global_dir)
        all_file_list = cal_detail_in_dir(global_dir)
        all_file_list = [Path(i) for i in all_file_list]
        all_file_list = [i for i in all_file_list if i.suffix in ['.pdf', '.docx']]
        print("-----in upload-dir, all_file_list:")
        print(all_file_list)
        all_trans_data = [transfile(i) for i in tqdm(all_file_list)]
        all_trans_data = [clean_text_data(i) for i in all_trans_data]
        all_trans_data = [i for i in all_trans_data if i.text_data.shape[0] > 0]
        print("all_trans_data:\n", all_trans_data)
        all_trans_data = [chunk_text4TransOutput(i) for i in all_trans_data]

        all_vector = [self.sv.encode_fun_plus(i.text_data['chunk_text'].tolist()) for i in all_trans_data]

        self.all_trans_data = all_trans_data
        self.all_vector = all_vector
        print("------update folder done")

if __name__ == "__main__":
    kl = KnowLedge(global_dir="data/data1",
                   gen_model_name_or_path="models/chatglm3-6b-32k",
                   sen_embedding_model_name_or_path="models/chinese-roberta-wwm-ext")
    while(True):
        query = input("请输入问题：")
        if query == "exit":
            exit()
        # res, data = kl.search_result("对话人尝试在哪里见面？")
        res, data = kl.search_result(query)
        print(res)
        print(data)
